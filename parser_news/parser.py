import requests
import os
from docx import Document, shared
from bs4 import BeautifulSoup


site = 'https://strana.ua/news'
domen = 'https://strana.ua'


def get_html(url):
    r = requests.get(url)
    return r.text


def get_total_pages(html):

    soup = BeautifulSoup(html, 'lxml')

    pages = soup.find('div', {'class': 'pagination'}).find_all('a')

    return pages[:4]


def get_url_pages(pages):

    five_pages = [site]

    for href in pages:
        five_pages.append(domen + href.get('href'))

    return five_pages


def get_page_links(html):

    soup = BeautifulSoup(html, 'lxml')

    titles = soup.find('div', {'class': 'lenta'}).find_all(
                                                      'div',
                                                      {'class': 'title'}
    )

    links = []

    for div in titles:
        a = div.find('a', {'class': 'article'}).get('href')
        link = domen + a
        links.append(link)

    return links


def get_all_links():

    all_links = []
    all_pages = get_url_pages(get_total_pages(get_html(site)))
    for page in all_pages:
        for link in get_page_links(get_html(page)):
            all_links.append(link)

    return all_links


def get_links_data(html):
    soup = BeautifulSoup(html, 'lxml')

    title = soup.find('h1', {'class': 'article'}).text.strip()
    object_news = soup.find('div', id='article-text').find_all('p')
    news = []
    for p in object_news:
        news.append(p.text)
    return title, news


def write_news(data):

    generate_file_name = 'News for OLEJA.docx'

    if generate_file_name in os.listdir(os.getcwd()):

        document = Document(generate_file_name)
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = shared.Pt(14)
        paragraph = document.add_paragraph()
        paragraph.add_run(data[0]).bold = True
        document.add_paragraph(data[1])
        document.save(generate_file_name)

    else:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = shared.Pt(14)
        paragraph = document.add_paragraph()
        paragraph.add_run(data[0]).bold = True
        document.add_paragraph(data[1])
        document.save(generate_file_name)


def main():
    for links in get_all_links():
        data = get_links_data(get_html(links))
        write_news(data)


if __name__ == '__main__':
    main()
